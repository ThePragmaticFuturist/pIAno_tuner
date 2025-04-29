import os
import requests
import pypdf
import docx # For reading .docx files
import csv   # For reading .csv files
import json
import time
import logging
import tempfile
from openai import OpenAI, RateLimitError, APIError
from tqdm import tqdm
from dotenv import load_dotenv
import re
from urllib.parse import urlparse

# --- Configuration ---
load_dotenv() # Load environment variables from .env file

SOURCE_FILE = "source_files.txt"  # File listing PDF/DOCX/CSV paths/URLs
OUTPUT_FILE = "training_data.jsonl" # Output JSON Lines file
CHUNK_SIZE = 1500   # Characters per chunk (adjust based on context length and desired granularity)
CHUNK_OVERLAP = 200 # Overlap between chunks to maintain context
MAX_RETRIES = 3     # Max retries for OpenAI API calls
RETRY_DELAY = 5     # Seconds to wait before retrying API calls
REQUEST_TIMEOUT = 120 # Timeout for downloading files in seconds

# --- OpenAI API Configuration ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = "gpt-3.5-turbo" # Or "gpt-4-turbo-preview", "gpt-4", etc.
# It's HIGHLY recommended to use a model that supports JSON mode if possible (like gpt-4-turbo-preview, gpt-3.5-turbo-1106+)

# --- Prompt Template for Data Generation ---
# (Keep the same template as before, instructing the LLM based on the provided text chunk)
GENERATION_PROMPT_TEMPLATE = """
Based *only* on the following text content extracted from a document, generate one relevant and specific instruction-response pair suitable for training an AI assistant.

The instruction should be a clear question or command derivable from the text.
The response should accurately answer the instruction using *only* information present in the provided text.
Do not add any information not present in the text.
Avoid generic instructions like "Summarize the text". Focus on specific concepts, definitions, explanations, processes, or data points mentioned.

Format the output strictly as a single JSON object containing two keys: "instruction" and "response".

Example Input Text (from PDF):
"Photosynthesis is the process used by plants, algae and cyanobacteria to convert light energy into chemical energy, through a process that uses water and carbon dioxide. Oxygen is released as a byproduct."

Example Output JSON:
{
  "instruction": "Explain the process of photosynthesis based on the provided text.",
  "response": "Photosynthesis is a process where plants, algae, and cyanobacteria use light energy, water, and carbon dioxide to create chemical energy, releasing oxygen."
}

Example Input Text (from CSV snippet):
"ProductID,ProductName,Price,Stock\n101,WidgetA,19.99,150\n102,GadgetB,25.50,75"

Example Output JSON:
{
  "instruction": "What is the price and stock level for ProductID 101 based on the data?",
  "response": "Based on the data, ProductID 101 (WidgetA) has a price of 19.99 and a stock level of 150."
}


Now, generate the JSON for the following text:

Text Content:
\"\"\"
{text_chunk}
\"\"\"

Output JSON:
"""

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Helper Functions ---

def read_source_list(filepath):
    """Reads the list of sources (PDF, DOCX, CSV) from a file."""
    if not os.path.exists(filepath):
        logging.error(f"Source file not found: {filepath}")
        return []
    with open(filepath, 'r', encoding='utf-8') as f:
        sources = [line.strip() for line in f if line.strip()]
    logging.info(f"Read {len(sources)} sources from {filepath}")
    return sources

def get_file_extension(filepath_or_url):
    """Gets the lowercase file extension from a path or URL."""
    try:
        path = urlparse(filepath_or_url).path
        return os.path.splitext(path)[1].lower()
    except Exception:
        return None

def download_file(url, target_path):
    """Downloads a file from a URL."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'} # Some servers block default requests user-agent
        response = requests.get(url, stream=True, timeout=REQUEST_TIMEOUT, headers=headers)
        response.raise_for_status() # Raise an exception for bad status codes
        with open(target_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        logging.info(f"Successfully downloaded file from {url} to {target_path}")
        return True
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to download {url}: {e}")
        return False
    except Exception as e:
        logging.error(f"An unexpected error occurred during download of {url}: {e}")
        return False

def clean_text(text):
    """Basic text cleaning."""
    if not text:
        return ""
    # Replace multiple whitespace characters with a single space
    text = re.sub(r'\s+', ' ', text)
    # Remove leading/trailing whitespace
    return text.strip()

def extract_text_from_pdf(pdf_path):
    """Extracts text content from a PDF file using pypdf."""
    text = ""
    try:
        reader = pypdf.PdfReader(pdf_path)
        num_pages = len(reader.pages)
        logging.info(f"Extracting text from {num_pages} pages in {os.path.basename(pdf_path)}...")
        for i, page in enumerate(tqdm(reader.pages, desc=f"Reading PDF {os.path.basename(pdf_path)}", unit="page", leave=False)):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n" # Add newline between pages
            except Exception as e:
                logging.warning(f"Could not extract text from page {i+1} in {os.path.basename(pdf_path)}: {e}")
        logging.info(f"Finished extracting text from {os.path.basename(pdf_path)}.")
        return clean_text(text)
    except Exception as e:
        logging.error(f"Failed to read PDF {pdf_path}: {e}")
        return None

def extract_text_from_docx(docx_path):
    """Extracts text content from a DOCX file."""
    text = ""
    try:
        logging.info(f"Extracting text from DOCX: {os.path.basename(docx_path)}...")
        document = docx.Document(docx_path)
        for para in tqdm(document.paragraphs, desc=f"Reading DOCX {os.path.basename(docx_path)}", unit="para", leave=False):
            text += para.text + "\n" # Add newline between paragraphs
        # Consider adding table extraction if needed:
        # for table in document.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             text += cell.text + "\t" # Use tab separation for table cells
        #         text += "\n"
        logging.info(f"Finished extracting text from {os.path.basename(docx_path)}.")
        return clean_text(text)
    except Exception as e:
        logging.error(f"Failed to read DOCX {docx_path}: {e}")
        return None

def extract_text_from_csv(csv_path):
    """Extracts text content from a CSV file, joining cells."""
    text = ""
    try:
        logging.info(f"Extracting text from CSV: {os.path.basename(csv_path)}...")
        with open(csv_path, 'r', newline='', encoding='utf-8', errors='ignore') as csvfile:
            # Sniff to detect dialect (delimiter, quoting)
            try:
                dialect = csv.Sniffer().sniff(csvfile.read(1024*10)) # Read more bytes for sniffing
                csvfile.seek(0) # Reset file pointer
                reader = csv.reader(csvfile, dialect)
            except csv.Error:
                logging.warning(f"Could not detect CSV dialect for {os.path.basename(csv_path)}. Falling back to standard comma delimiter.")
                csvfile.seek(0)
                reader = csv.reader(csvfile) # Default comma delimiter

            header = next(reader, None) # Read header row
            if header:
                text += ", ".join(header) + "\n" # Add header representation

            for row in tqdm(reader, desc=f"Reading CSV {os.path.basename(csv_path)}", unit="row", leave=False):
                # Join cells in the row with a comma and space
                text += ", ".join(row) + "\n"
        logging.info(f"Finished extracting text from {os.path.basename(csv_path)}.")
        return clean_text(text) # Apply basic cleaning
    except Exception as e:
        logging.error(f"Failed to read CSV {csv_path}: {e}")
        return None

def chunk_text(text, chunk_size, chunk_overlap):
    """Splits text into overlapping chunks."""
    # (Keep the same chunking logic as before)
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = start + chunk_size
        chunks.append(text[start:end])
        next_start = start + chunk_size - chunk_overlap
        # Prevent infinite loop if overlap >= chunk_size or step is too small
        if next_start <= start:
             start += 1 # Move forward by at least one character
        else:
             start = next_start

    logging.info(f"Split text into {len(chunks)} chunks.")
    return chunks


def generate_training_pairs(text_chunk, client, model, prompt_template):
    """Generates instruction/response pairs using the OpenAI API."""
    # (Keep the same generation logic as before)
    prompt = prompt_template.format(text_chunk=text_chunk)
    retries = 0
    while retries < MAX_RETRIES:
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an AI assistant tasked with creating training data. Follow the user's instructions precisely and output only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                # For models supporting JSON mode (like gpt-4-turbo-preview, gpt-3.5-turbo-1106+):
                response_format={ "type": "json_object" },
                temperature=0.3, # Lower temperature for more deterministic output
                max_tokens=350  # Slightly increase if responses from CSV/DOCX might be longer
            )
            content = response.choices[0].message.content
            try:
                data = json.loads(content)
                if isinstance(data, dict) and "instruction" in data and "response" in data:
                    if data["instruction"] and data["response"] and len(data["instruction"]) > 5 and len(data["response"]) > 5:
                         if "based on the provided text" not in data["instruction"].lower() and \
                            "based on the text" not in data["response"].lower() and \
                            "not mentioned in the text" not in data["response"].lower():
                                return data
                         else:
                            logging.warning(f"Generated pair might contain placeholder/refusal text: {data}")
                            return None
                    else:
                        logging.warning(f"Generated JSON missing fields or fields too short: {content}")
                        return None
                else:
                    logging.warning(f"Generated content is not a valid JSON object with required keys: {content}")
                    return None
            except json.JSONDecodeError:
                logging.warning(f"Failed to decode JSON from LLM response: {content}")
                return None

        except RateLimitError as e:
            logging.warning(f"Rate limit exceeded. Retrying in {RETRY_DELAY} seconds... ({retries+1}/{MAX_RETRIES})")
            retries += 1
            time.sleep(RETRY_DELAY)
        except APIError as e:
            logging.error(f"OpenAI API error: {e}. Retrying in {RETRY_DELAY} seconds... ({retries+1}/{MAX_RETRIES})")
            retries += 1
            time.sleep(RETRY_DELAY)
        except Exception as e:
            logging.error(f"An unexpected error occurred during API call: {e}")
            return None # Don't retry on unexpected errors
    logging.error(f"Max retries reached for API call for chunk starting with: {text_chunk[:100]}...")
    return None


def format_for_finetuning(instruction, response):
    """Formats the instruction/response pair into the desired final structure."""
    # (Keep the same formatting logic as before)
    formatted_text = f"Below is an instruction that describes a task. Write a response that appropriately completes the request.\n\n### Instruction:\n{instruction}\n\n### Response:\n{response}"
    return json.dumps({"text": formatted_text})


# --- Main Execution ---

def main():
    if not OPENAI_API_KEY:
        logging.error("OPENAI_API_KEY environment variable not set. Please create a .env file or set the variable.")
        return

    try:
        client = OpenAI() # Initializes with API key from environment variable
        # Test connection (optional but good practice)
        client.models.list()
        logging.info("Successfully connected to OpenAI API.")
    except Exception as e:
        logging.error(f"Failed to initialize or connect to OpenAI API: {e}")
        return


    sources = read_source_list(SOURCE_FILE)
    if not sources:
        return

    generated_count = 0
    processed_files = 0
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as outfile:
        for source in tqdm(sources, desc="Processing Sources"):
            source_path = None
            temp_file_info = None # To store {'path': path, 'ext': ext} for temp files
            source_ext = get_file_extension(source)
            is_url = source.startswith("http://") or source.startswith("https://")

            if is_url:
                if source_ext not in ['.pdf', '.docx', '.csv']:
                    logging.warning(f"Skipping URL with unsupported extension: {source}")
                    continue
                # Download URL to a temporary file
                # Create temp file with the correct suffix if possible
                suffix = source_ext if source_ext else None
                try:
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                    if download_file(source, temp_file.name):
                         temp_file_info = {'path': temp_file.name, 'ext': source_ext}
                         source_path = temp_file.name # Use temp path for processing
                    else:
                         temp_file.close()
                         os.unlink(temp_file.name) # Clean up temp file on download failure
                         continue # Skip to next source
                except Exception as e:
                     logging.error(f"Error creating/handling temp file for {source}: {e}")
                     if 'temp_file' in locals() and temp_file:
                          try:
                              temp_file.close()
                              if os.path.exists(temp_file.name):
                                   os.unlink(temp_file.name)
                          except Exception: pass
                     continue

            elif os.path.exists(source):
                # It's a local file path
                if source_ext not in ['.pdf', '.docx', '.csv']:
                    logging.warning(f"Skipping local file with unsupported extension: {source}")
                    continue
                source_path = source
            else:
                logging.warning(f"Source not found or invalid URL/Path: {source}")
                continue

            # Determine which extractor to use based on the definitive extension
            # For URLs, we rely on the extension derived from the URL itself.
            # For local files, we use the actual file extension.
            file_ext_to_use = temp_file_info['ext'] if temp_file_info else source_ext

            extracted_text = None
            if source_path and file_ext_to_use:
                logging.info(f"Processing file ({file_ext_to_use}): {os.path.basename(source_path)}")
                try:
                    if file_ext_to_use == '.pdf':
                        extracted_text = extract_text_from_pdf(source_path)
                    elif file_ext_to_use == '.docx':
                        extracted_text = extract_text_from_docx(source_path)
                    elif file_ext_to_use == '.csv':
                        extracted_text = extract_text_from_csv(source_path)
                    processed_files += 1
                except Exception as e:
                    logging.error(f"Unhandled error during text extraction for {source_path}: {e}")
                    extracted_text = None # Ensure it's None on error


            # Process extracted text (chunking and generation)
            if extracted_text:
                text_chunks = chunk_text(extracted_text, CHUNK_SIZE, CHUNK_OVERLAP)
                logging.info(f"Generating training data from {len(text_chunks)} chunks for {os.path.basename(source_path)}...")
                chunk_counter = 0
                for chunk in tqdm(text_chunks, desc=f"Generating pairs for {os.path.basename(source_path)}", leave=False):
                    if len(chunk.strip()) < 50: # Skip very short chunks
                        continue
                    generated_data = generate_training_pairs(chunk, client, OPENAI_MODEL, GENERATION_PROMPT_TEMPLATE)
                    if generated_data:
                        finetuning_entry = format_for_finetuning(generated_data["instruction"], generated_data["response"])
                        outfile.write(finetuning_entry + "\n")
                        generated_count += 1
                        chunk_counter += 1
                        # Optional: Add a small delay to avoid hitting rate limits too quickly
                        time.sleep(0.2) # Reduce sleep time slightly maybe?
                logging.info(f"Generated {chunk_counter} pairs from {os.path.basename(source_path)}")
            elif source_path:
                logging.warning(f"No text could be extracted from {source_path}.")


            # Clean up temporary file if it was created
            if temp_file_info:
                try:
                    # temp_file object might be out of scope, use path directly
                    if os.path.exists(temp_file_info['path']):
                        os.unlink(temp_file_info['path'])
                except Exception as e:
                    logging.warning(f"Could not clean up temp file {temp_file_info['path']}: {e}")


    logging.info(f"--- Processing Complete ---")
    logging.info(f"Processed {processed_files} files.")
    logging.info(f"Successfully generated {generated_count} instruction/response pairs.")
    logging.info(f"Training data saved to: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
